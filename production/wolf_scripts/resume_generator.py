#!/usr/bin/env python3
"""
ATS-Optimized Resume Generation Engine
Parses job postings, queries pgai for relevant experience, generates ATS-friendly resumes.

Usage:
    python resume_generator.py --mode generate --job-posting <text> --output <dir>
    python resume_generator.py --mode track --company <name> --job-id <id> --resume-id <id>
    python resume_generator.py --mode list-memories --query <text>
"""

import os
import sys
import json
import re
import argparse
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Optional
from pathlib import Path
import psycopg2
from psycopg2.extras import RealDictCursor
import threading
from concurrent.futures import ThreadPoolExecutor

# PDF and DOCX generation
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
import requests

# Embedding and LLM
try:
    import ollama
except ImportError:
    print("Warning: ollama not available. Install with: pip install ollama")
    ollama = None

# ============================================================================
# CONFIGURATION
# ============================================================================

PG_CONFIG = {
    "host": "localhost",
    "port": 5433,
    "database": "wolf_logic",
    "user": "wolf",
    "password": "wolflogic2024"
}

EMBEDDING_MODEL = "nomic-embed-text:v1.5"

# ATS-Friendly formatting
ATS_FONTS = ["Arial", "Calibri", "Times New Roman", "Helvetica"]
ATS_FONT_SIZE = 11
ATS_LINE_SPACING = 1.15

# Resume sections (order matters for ATS)
RESUME_SECTIONS = [
    "PROFESSIONAL SUMMARY",
    "SKILLS",
    "PROFESSIONAL EXPERIENCE",
    "EDUCATION",
    "CERTIFICATIONS",
    "ADDITIONAL"
]

print_lock = threading.Lock()

def safe_print(msg):
    """Thread-safe printing"""
    with print_lock:
        print(f"[{datetime.now().strftime('%H:%M:%S')}] {msg}")


# ============================================================================
# DATABASE OPERATIONS
# ============================================================================

class PgaiDB:
    """PostgreSQL pgai connector for resume data"""

    def __init__(self, config: Dict):
        self.config = config
        self.conn = None

    def connect(self):
        """Establish database connection"""
        try:
            self.conn = psycopg2.connect(**self.config)
            safe_print(f"Connected to {self.config['database']}")
        except psycopg2.Error as e:
            safe_print(f"Database connection failed: {e}")
            raise

    def close(self):
        """Close database connection"""
        if self.conn:
            self.conn.close()

    def query_memories(self, search_query: str, limit: int = 10) -> List[Dict]:
        """Semantic search through pgai memories using embeddings"""
        if not self.conn:
            self.connect()

        try:
            # Generate embedding for search query
            if ollama:
                embedding_response = ollama.embeddings(
                    model=EMBEDDING_MODEL,
                    prompt=search_query
                )
                query_embedding = embedding_response["embedding"]
            else:
                safe_print(f"Warning: ollama not available, using keyword search")
                return self.keyword_search(search_query, limit)

            # Vector similarity search in pgai
            cur = self.conn.cursor(cursor_factory=RealDictCursor)

            # Using pgvector similarity search
            query = """
                SELECT
                    id,
                    content,
                    namespace,
                    created_at,
                    metadata,
                    1 - (embedding <=> %s) as similarity
                FROM memories_embedding
                WHERE namespace NOT IN ('internal', 'system')
                ORDER BY embedding <=> %s
                LIMIT %s
            """

            cur.execute(query, (json.dumps(query_embedding), json.dumps(query_embedding), limit))
            results = cur.fetchall()
            cur.close()

            return results

        except Exception as e:
            safe_print(f"Query memories error: {e}")
            return self.keyword_search(search_query, limit)

    def keyword_search(self, search_query: str, limit: int = 10) -> List[Dict]:
        """Fallback keyword search in memories"""
        try:
            cur = self.conn.cursor(cursor_factory=RealDictCursor)

            keywords = search_query.split()
            search_pattern = " & ".join(keywords)

            query = """
                SELECT
                    id,
                    content,
                    namespace,
                    created_at,
                    metadata
                FROM memories
                WHERE to_tsvector('english', content) @@ to_tsquery('english', %s)
                LIMIT %s
            """

            cur.execute(query, (search_pattern, limit))
            results = cur.fetchall()
            cur.close()

            return results

        except Exception as e:
            safe_print(f"Keyword search error: {e}")
            return []

    def get_user_profile(self) -> Optional[Dict]:
        """Get user profile information from memories"""
        try:
            cur = self.conn.cursor(cursor_factory=RealDictCursor)

            cur.execute("""
                SELECT content, metadata
                FROM memories
                WHERE namespace = 'profile' OR namespace = 'personal'
                ORDER BY created_at DESC
                LIMIT 1
            """)

            result = cur.fetchone()
            cur.close()

            return result
        except Exception as e:
            safe_print(f"Get profile error: {e}")
            return None

    def create_resume_table(self):
        """Create resume tracking table if not exists"""
        try:
            cur = self.conn.cursor()

            cur.execute("""
                CREATE TABLE IF NOT EXISTS resume_applications (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER,
                    company_name VARCHAR(255),
                    job_title VARCHAR(255),
                    job_posting TEXT,
                    resume_id VARCHAR(255) UNIQUE,
                    resume_format VARCHAR(50),
                    extraction_keywords JSONB,
                    matched_skills JSONB,
                    years_required INTEGER,
                    sent_date TIMESTAMP DEFAULT NOW(),
                    created_at TIMESTAMP DEFAULT NOW()
                )
            """)

            self.conn.commit()
            safe_print("Resume tracking table created/verified")

        except psycopg2.Error as e:
            safe_print(f"Table creation error: {e}")
            self.conn.rollback()

    def save_resume_application(self, company: str, job_title: str, job_posting: str,
                               resume_id: str, keywords: Dict, matched_skills: List,
                               years_required: int) -> bool:
        """Save resume application tracking record"""
        try:
            cur = self.conn.cursor()

            cur.execute("""
                INSERT INTO resume_applications
                (company_name, job_title, job_posting, resume_id, extraction_keywords,
                 matched_skills, years_required, resume_format)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                company,
                job_title,
                job_posting,
                resume_id,
                json.dumps(keywords),
                json.dumps(matched_skills),
                years_required,
                "DOCX/PDF"
            ))

            self.conn.commit()
            safe_print(f"Saved resume application: {resume_id} -> {company}")
            return True

        except psycopg2.Error as e:
            safe_print(f"Save application error: {e}")
            self.conn.rollback()
            return False


# ============================================================================
# JOB POSTING PARSING
# ============================================================================

class JobPostingParser:
    """Extract structured data from job postings"""

    # Common technical skills patterns
    SKILL_PATTERNS = {
        'languages': [
            r'\bPython\b', r'\bJava\b', r'\bJavaScript\b', r'\bTypeScript\b',
            r'\bC\+\+\b', r'\bC#\b', r'\bRust\b', r'\bGo\b', r'\bRuby\b',
            r'\bPHP\b', r'\bSwift\b', r'\bKotlin\b', r'\bSQL\b', r'\bR\b'
        ],
        'frameworks': [
            r'\bReact\b', r'\bAngular\b', r'\bVue\b', r'\bDjango\b',
            r'\bFastAPI\b', r'\bFlask\b', r'\bSpring\b', r'\bASP\.NET\b',
            r'\bNode\.js\b', r'\bExpress\b', r'\bNuxt\b'
        ],
        'databases': [
            r'\bPostgreSQL\b', r'\bMySQL\b', r'\bMongoDB\b', r'\bRedis\b',
            r'\bElasticsearch\b', r'\bCassandra\b', r'\bDynamoDB\b',
            r'\bFirebase\b', r'\bSQLite\b'
        ],
        'cloud': [
            r'\bAWS\b', r'\bAzure\b', r'\bGoogle Cloud\b', r'\bGCP\b',
            r'\bDocker\b', r'\bKubernetes\b', r'\bCI/CD\b'
        ],
        'tools': [
            r'\bGit\b', r'\bDocker\b', r'\bJenkins\b', r'\bGitHub\b',
            r'\bGitLab\b', r'\bJira\b', r'\bLinux\b'
        ]
    }

    def __init__(self, job_posting_text: str):
        self.text = job_posting_text
        self.extracted = {}

    def parse(self) -> Dict:
        """Extract key information from job posting"""
        self.extracted = {
            'job_title': self._extract_job_title(),
            'company': self._extract_company(),
            'required_skills': self._extract_required_skills(),
            'years_experience': self._extract_years_experience(),
            'key_technologies': self._extract_key_technologies(),
            'responsibilities': self._extract_responsibilities(),
            'nice_to_have': self._extract_nice_to_have(),
            'raw_keywords': self._extract_all_keywords(),
        }

        safe_print(f"Extracted job title: {self.extracted.get('job_title', 'Unknown')}")
        safe_print(f"Found {len(self.extracted.get('required_skills', []))} required skills")

        return self.extracted

    def _extract_job_title(self) -> str:
        """Extract job title"""
        lines = self.text.split('\n')

        # Look for common title patterns
        for line in lines[:20]:  # Check first 20 lines
            if any(keyword in line.lower() for keyword in ['job title', 'position', 'role']):
                return line.split(':', 1)[-1].strip()[:100]

        # Fallback to first substantial line
        for line in lines:
            if len(line.strip()) > 10 and len(line.strip()) < 200:
                return line.strip()[:100]

        return "Unknown Position"

    def _extract_company(self) -> str:
        """Extract company name"""
        lines = self.text.split('\n')

        for line in lines[:15]:
            if any(keyword in line.lower() for keyword in ['company', 'about us', 'organization']):
                return line.split(':', 1)[-1].strip()[:100]

        return "Unknown Company"

    def _extract_years_experience(self) -> int:
        """Extract years of experience required"""
        # Look for patterns like "5+ years", "3-5 years", etc.
        pattern = r'(\d+)\s*(?:\+|-\d+)?\s*(?:year|yr)s?\s+(?:of\s+)?(?:experience|exp)'

        matches = re.finditer(pattern, self.text, re.IGNORECASE)

        for match in matches:
            try:
                return int(match.group(1))
            except:
                pass

        return 0  # Default if not found

    def _extract_required_skills(self) -> List[str]:
        """Extract required skills from job posting"""
        skills = set()

        # Look for skills section
        skills_section = self._find_section(['required skills', 'must have', 'requirements'])

        if skills_section:
            text_to_search = skills_section
        else:
            text_to_search = self.text

        # Search for all skill patterns
        for category, patterns in self.SKILL_PATTERNS.items():
            for pattern in patterns:
                matches = re.finditer(pattern, text_to_search, re.IGNORECASE)
                for match in matches:
                    skills.add(match.group())

        # Extract skills from bullet points
        bullet_pattern = r'(?:^|\n)\s*[-•*]\s+([^\n]+)'
        for match in re.finditer(bullet_pattern, text_to_search):
            line = match.group(1)
            # Check if line contains skill-like keywords
            if any(skill_word in line.lower() for skill_word in
                   ['experience', 'knowledge', 'proficiency', 'expertise', 'understand']):
                skills.add(line.strip()[:100])

        return sorted(list(skills))

    def _extract_key_technologies(self) -> List[str]:
        """Extract key technologies mentioned"""
        tech = set()

        # Technology-specific patterns
        tech_patterns = [
            r'\b(?:AWS|Azure|GCP|Google Cloud|Kubernetes|Docker)\b',
            r'\b(?:React|Vue|Angular|Svelte)\b',
            r'\b(?:Python|Java|Go|Rust|JavaScript|TypeScript)\b',
            r'\b(?:PostgreSQL|MongoDB|MySQL|Redis)\b',
        ]

        for pattern in tech_patterns:
            matches = re.finditer(pattern, self.text, re.IGNORECASE)
            for match in matches:
                tech.add(match.group())

        return sorted(list(tech))

    def _extract_responsibilities(self) -> List[str]:
        """Extract key responsibilities"""
        resp = []

        resp_section = self._find_section(['responsibilities', 'you will', 'what you', 'your role'])

        if resp_section:
            # Extract bullet points
            bullets = re.findall(r'(?:^|\n)\s*[-•*]\s+([^\n]+)', resp_section)
            resp = [b.strip()[:150] for b in bullets[:5]]  # Top 5

        return resp

    def _extract_nice_to_have(self) -> List[str]:
        """Extract nice-to-have skills"""
        nice = []

        nice_section = self._find_section(['nice to have', 'preferred', 'bonus'])

        if nice_section:
            bullets = re.findall(r'(?:^|\n)\s*[-•*]\s+([^\n]+)', nice_section)
            nice = [b.strip()[:100] for b in bullets[:5]]

        return nice

    def _find_section(self, keywords: List[str]) -> Optional[str]:
        """Find section by keywords"""
        text_lower = self.text.lower()

        for keyword in keywords:
            idx = text_lower.find(keyword)
            if idx != -1:
                # Find section start and end
                start = idx
                # End at next section header or 2000 chars
                end = start + 2000

                section = self.text[start:end]

                # Try to find end of section (next "Section:" pattern)
                next_section = re.search(r'\n[A-Z][A-Z\s]+:', section[len(keyword):])
                if next_section:
                    end = start + len(keyword) + next_section.start()

                return self.text[start:end]

        return None

    def _extract_all_keywords(self) -> List[str]:
        """Extract all meaningful keywords from posting"""
        # Remove common words
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'is', 'are', 'be', 'been', 'being', 'have', 'has', 'had',
            'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might'
        }

        words = re.findall(r'\b[a-z]+\b', self.text.lower())
        keywords = [w for w in words if w not in stop_words and len(w) > 3]

        # Count frequency
        from collections import Counter
        freq = Counter(keywords)

        # Return top keywords
        return [word for word, _ in freq.most_common(30)]


# ============================================================================
# RESUME BUILDER
# ============================================================================

class ResumeBuilder:
    """Generate ATS-optimized resumes in DOCX and PDF"""

    def __init__(self, user_data: Dict, job_analysis: Dict, pgai: PgaiDB):
        self.user_data = user_data
        self.job_analysis = job_analysis
        self.pgai = pgai
        self.resume_content = {}

    def build(self, skills_to_highlight: List[str]) -> Dict:
        """Build complete resume content"""
        safe_print("Building resume content...")

        self.resume_content = {
            'header': self._build_header(),
            'summary': self._build_summary(skills_to_highlight),
            'skills': self._build_skills(skills_to_highlight),
            'experience': self._build_experience(skills_to_highlight),
            'education': self._build_education(),
            'certifications': self._build_certifications(),
        }

        return self.resume_content

    def _build_header(self) -> Dict:
        """Build resume header with contact info"""
        return {
            'name': self.user_data.get('name', 'Your Name'),
            'title': self.user_data.get('title', 'Professional'),
            'email': self.user_data.get('email', 'email@example.com'),
            'phone': self.user_data.get('phone', '+1 (555) 000-0000'),
            'location': self.user_data.get('location', 'City, State'),
            'linkedin': self.user_data.get('linkedin', ''),
            'github': self.user_data.get('github', ''),
        }

    def _build_summary(self, job_skills: List[str]) -> str:
        """Build professional summary tailored to job"""
        years = self.job_analysis.get('years_experience', 5)
        job_title = self.job_analysis.get('job_title', 'Professional Role')
        top_skills = job_skills[:3]

        # Create tailored summary
        summary = f"Results-driven professional with {years}+ years of experience in {', '.join(top_skills)}. "
        summary += "Proven track record of delivering high-quality solutions and driving technical excellence. "
        summary += f"Seeking a {job_title} position to leverage expertise and contribute to organizational growth."

        return summary

    def _build_skills(self, job_skills: List[str]) -> Dict:
        """Build skills section matching job requirements"""
        # Organize by category
        skills_by_category = {
            'Languages': [],
            'Frameworks & Libraries': [],
            'Databases & Tools': [],
            'Cloud & DevOps': [],
        }

        # Retrieve skill details from memories
        all_user_skills = self.pgai.query_memories("technical skills experience", limit=10)

        # Categorize job skills
        for skill in job_skills[:15]:  # Limit to top 15
            if any(lang in skill.lower() for lang in ['python', 'java', 'javascript', 'c++', 'go', 'rust']):
                skills_by_category['Languages'].append(skill)
            elif any(fw in skill.lower() for fw in ['react', 'vue', 'django', 'spring', 'node']):
                skills_by_category['Frameworks & Libraries'].append(skill)
            elif any(db in skill.lower() for db in ['sql', 'mongo', 'redis', 'elastic', 'docker']):
                skills_by_category['Databases & Tools'].append(skill)
            elif any(cloud in skill.lower() for cloud in ['aws', 'azure', 'gcp', 'kubernetes']):
                skills_by_category['Cloud & DevOps'].append(skill)
            else:
                skills_by_category['Languages'].append(skill)

        return skills_by_category

    def _build_experience(self, job_skills: List[str]) -> List[Dict]:
        """Build experience section with memory-driven content"""
        experiences = []

        # Query memories for relevant experiences
        for skill in job_skills[:5]:  # Top 5 skills
            results = self.pgai.query_memories(f"professional experience with {skill}", limit=3)

            for result in results:
                if isinstance(result, dict):
                    content = result.get('content', '')

                    # Parse experience bullet points
                    experience = {
                        'company': 'Your Company',
                        'title': f"{skill} Professional",
                        'duration': '2022 - Present',
                        'description': content[:200],
                        'bullets': [
                            f"Utilized {skill} to deliver high-quality solutions",
                            "Improved system performance and reliability",
                            "Collaborated with cross-functional teams",
                        ]
                    }

                    experiences.append(experience)

        # Default if no memories found
        if not experiences:
            experiences.append({
                'company': 'Previous Organization',
                'title': 'Technical Professional',
                'duration': '2022 - Present',
                'description': f"Developed solutions using {', '.join(job_skills[:3])}",
                'bullets': [
                    f"Specialized in {job_skills[0] if job_skills else 'technical solutions'}",
                    "Delivered projects on time and within budget",
                    "Mentored junior team members",
                ]
            })

        return experiences[:4]  # Limit to 4 positions

    def _build_education(self) -> List[Dict]:
        """Build education section"""
        return [
            {
                'degree': self.user_data.get('degree', 'Bachelor of Science'),
                'field': self.user_data.get('field', 'Computer Science'),
                'school': self.user_data.get('school', 'University Name'),
                'graduation': self.user_data.get('graduation', '2020'),
            }
        ]

    def _build_certifications(self) -> List[str]:
        """Build certifications section"""
        return self.user_data.get('certifications', [
            'AWS Certified Solutions Architect',
            'Professional Scrum Master',
        ])

    def to_docx(self, output_path: str) -> bool:
        """Generate DOCX resume"""
        try:
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(ATS_FONT_SIZE)

            # ===== HEADER =====
            header = self.resume_content['header']

            # Name (largest)
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(header['name'].upper())
            name_run.font.size = Pt(14)
            name_run.font.bold = True
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Title
            title_para = doc.add_paragraph(header['title'])
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contact info
            contact = f"{header['email']} | {header['phone']} | {header['location']}"
            contact_para = doc.add_paragraph(contact)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacing

            # ===== PROFESSIONAL SUMMARY =====
            doc.add_paragraph('PROFESSIONAL SUMMARY', style='Heading 1').runs[0].font.size = Pt(11)
            doc.add_paragraph(self.resume_content['summary'])
            doc.add_paragraph()  # Spacing

            # ===== SKILLS =====
            doc.add_paragraph('SKILLS', style='Heading 1').runs[0].font.size = Pt(11)

            for category, skills in self.resume_content['skills'].items():
                if skills:
                    skill_text = f"{category}: {', '.join(skills[:8])}"  # Limit to 8 per category
                    doc.add_paragraph(skill_text, style='List Bullet')

            doc.add_paragraph()  # Spacing

            # ===== PROFESSIONAL EXPERIENCE =====
            doc.add_paragraph('PROFESSIONAL EXPERIENCE', style='Heading 1').runs[0].font.size = Pt(11)

            for exp in self.resume_content['experience']:
                # Position title and company
                exp_header = f"{exp['title']} | {exp['company']}"
                exp_para = doc.add_paragraph(exp_header)
                exp_para.runs[0].font.bold = True

                # Duration
                doc.add_paragraph(exp['duration'], style='List Number')

                # Bullets
                for bullet in exp['bullets']:
                    doc.add_paragraph(bullet, style='List Bullet')

                doc.add_paragraph()  # Spacing

            # ===== EDUCATION =====
            doc.add_paragraph('EDUCATION', style='Heading 1').runs[0].font.size = Pt(11)

            for edu in self.resume_content['education']:
                edu_para = doc.add_paragraph(f"{edu['degree']} in {edu['field']}")
                edu_para.runs[0].font.bold = True
                doc.add_paragraph(f"{edu['school']} | Graduated: {edu['graduation']}")

            doc.add_paragraph()  # Spacing

            # ===== CERTIFICATIONS =====
            if self.resume_content['certifications']:
                doc.add_paragraph('CERTIFICATIONS', style='Heading 1').runs[0].font.size = Pt(11)
                for cert in self.resume_content['certifications']:
                    doc.add_paragraph(cert, style='List Bullet')

            # Save document
            doc.save(output_path)
            safe_print(f"DOCX resume saved: {output_path}")
            return True

        except Exception as e:
            safe_print(f"DOCX generation error: {e}")
            return False

    def to_pdf(self, output_path: str) -> bool:
        """Generate PDF resume"""
        try:
            # PDF setup
            pdf_file = SimpleDocTemplate(output_path, pagesize=letter,
                                        rightMargin=0.5*inch, leftMargin=0.5*inch,
                                        topMargin=0.5*inch, bottomMargin=0.5*inch)

            story = []
            styles = getSampleStyleSheet()

            # Define ATS-friendly styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Normal'],
                fontSize=14,
                textColor=colors.black,
                spaceAfter=6,
                alignment=TA_LEFT,
                fontName='Helvetica-Bold'
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Normal'],
                fontSize=11,
                textColor=colors.black,
                spaceAfter=4,
                spaceBefore=8,
                fontName='Helvetica-Bold'
            )

            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=10,
                alignment=TA_JUSTIFY,
                spaceAfter=3
            )

            header = self.resume_content['header']

            # Title
            story.append(Paragraph(header['name'].upper(), title_style))
            story.append(Paragraph(header['title'], styles['Normal']))

            # Contact
            contact = f"{header['email']} | {header['phone']} | {header['location']}"
            story.append(Paragraph(contact, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))

            # Summary
            story.append(Paragraph('PROFESSIONAL SUMMARY', heading_style))
            story.append(Paragraph(self.resume_content['summary'], body_style))
            story.append(Spacer(1, 0.05*inch))

            # Skills
            story.append(Paragraph('SKILLS', heading_style))
            for category, skills in self.resume_content['skills'].items():
                if skills:
                    skill_text = f"<b>{category}:</b> {', '.join(skills[:8])}"
                    story.append(Paragraph(skill_text, body_style))
            story.append(Spacer(1, 0.05*inch))

            # Experience
            story.append(Paragraph('PROFESSIONAL EXPERIENCE', heading_style))
            for exp in self.resume_content['experience']:
                exp_text = f"<b>{exp['title']} | {exp['company']}</b>"
                story.append(Paragraph(exp_text, body_style))
                story.append(Paragraph(exp['duration'], styles['Normal']))

                for bullet in exp['bullets']:
                    bullet_text = f"• {bullet}"
                    story.append(Paragraph(bullet_text, body_style))
                story.append(Spacer(1, 0.03*inch))

            # Education
            story.append(Paragraph('EDUCATION', heading_style))
            for edu in self.resume_content['education']:
                edu_text = f"<b>{edu['degree']} in {edu['field']}</b>"
                story.append(Paragraph(edu_text, body_style))
                story.append(Paragraph(f"{edu['school']} | Graduated: {edu['graduation']}", styles['Normal']))

            # Build PDF
            pdf_file.build(story)
            safe_print(f"PDF resume saved: {output_path}")
            return True

        except Exception as e:
            safe_print(f"PDF generation error: {e}")
            return False


# ============================================================================
# SKILL MATCHING ENGINE
# ============================================================================

class SkillMatcher:
    """Match user skills with job requirements"""

    def __init__(self, pgai: PgaiDB):
        self.pgai = pgai

    def match_skills(self, required_skills: List[str], nice_to_have: List[str]) -> Tuple[List[str], Dict]:
        """Match user skills with job requirements"""

        matched_skills = []
        match_scores = {}

        # Check user memories for skill matches
        for skill in required_skills:
            # Query for experience with this skill
            memories = self.pgai.query_memories(f"experience with {skill}", limit=3)

            if memories:
                matched_skills.append(skill)
                match_scores[skill] = {
                    'confidence': 'high',
                    'evidence_count': len(memories)
                }
                safe_print(f"  Matched required skill: {skill}")
            else:
                match_scores[skill] = {
                    'confidence': 'low',
                    'evidence_count': 0
                }

        # Check nice-to-have
        for skill in nice_to_have:
            memories = self.pgai.query_memories(f"experience with {skill}", limit=2)
            if memories:
                matched_skills.append(skill)
                match_scores[skill] = {
                    'confidence': 'medium',
                    'evidence_count': len(memories)
                }
                safe_print(f"  Matched nice-to-have skill: {skill}")

        return matched_skills, match_scores

    def prioritize_skills(self, all_skills: List[str], max_count: int = 15) -> List[str]:
        """Prioritize skills for resume"""
        # Skills should be in order: most relevant first

        # Common ranking order
        priority_order = {
            'Python': 1, 'Java': 1, 'JavaScript': 1, 'TypeScript': 1,
            'React': 2, 'Django': 2, 'FastAPI': 2,
            'PostgreSQL': 3, 'MongoDB': 3, 'Redis': 3,
            'AWS': 4, 'Docker': 4, 'Kubernetes': 4,
            'Git': 5, 'CI/CD': 5,
        }

        # Sort by priority
        prioritized = sorted(
            all_skills,
            key=lambda x: (priority_order.get(x, 10), all_skills.index(x))
        )

        return prioritized[:max_count]


# ============================================================================
# MAIN ORCHESTRATION
# ============================================================================

class ResumeGenerationEngine:
    """Main orchestration engine"""

    def __init__(self):
        self.pgai = PgaiDB(PG_CONFIG)
        self.pgai.connect()
        self.pgai.create_resume_table()

    def generate(self, job_posting_text: str, output_dir: str,
                company_name: str = "Unknown Company") -> Dict:
        """Main generation pipeline"""

        safe_print("="*70)
        safe_print("RESUME GENERATION ENGINE - STARTING")
        safe_print("="*70)

        results = {
            'success': False,
            'job_analysis': {},
            'matched_skills': [],
            'output_files': {},
            'resume_id': None,
            'tracking_record': None,
        }

        try:
            # Step 1: Parse job posting
            safe_print("\n[1/6] Parsing job posting...")
            parser = JobPostingParser(job_posting_text)
            job_analysis = parser.parse()
            results['job_analysis'] = job_analysis

            # Step 2: Match skills
            safe_print("\n[2/6] Matching skills with pgai memories...")
            matcher = SkillMatcher(self.pgai)

            required_skills = job_analysis.get('required_skills', [])
            nice_to_have = job_analysis.get('nice_to_have', [])

            matched, match_scores = matcher.match_skills(required_skills, nice_to_have)
            prioritized_skills = matcher.prioritize_skills(matched + required_skills)

            results['matched_skills'] = prioritized_skills
            safe_print(f"  Matched {len(prioritized_skills)} skills for highlighting")

            # Step 3: Get user profile
            safe_print("\n[3/6] Loading user profile...")
            user_profile = self.pgai.get_user_profile()

            user_data = {
                'name': 'Your Name',
                'title': 'Software Engineer',
                'email': 'your.email@example.com',
                'phone': '+1 (555) 123-4567',
                'location': 'San Francisco, CA',
                'linkedin': 'linkedin.com/in/yourname',
                'github': 'github.com/yourname',
                'degree': 'Bachelor of Science',
                'field': 'Computer Science',
                'school': 'University Name',
                'graduation': '2020',
                'certifications': ['AWS Certified Solutions Architect', 'Professional Scrum Master'],
            }

            if user_profile and isinstance(user_profile, dict):
                # Parse profile data if available
                profile_content = user_profile.get('content', '')
                safe_print(f"  Found user profile data")

            # Step 4: Build resume content
            safe_print("\n[4/6] Building ATS-optimized resume content...")
            builder = ResumeBuilder(user_data, job_analysis, self.pgai)
            resume_content = builder.build(prioritized_skills)

            # Step 5: Generate output files
            safe_print("\n[5/6] Generating resume files...")

            # Create output directory
            Path(output_dir).mkdir(parents=True, exist_ok=True)

            # Generate unique resume ID
            import uuid
            resume_id = f"resume_{company_name.replace(' ', '_')}_{uuid.uuid4().hex[:8]}"

            # DOCX output
            docx_path = os.path.join(output_dir, f"{resume_id}.docx")
            docx_success = builder.to_docx(docx_path)
            if docx_success:
                results['output_files']['docx'] = docx_path

            # PDF output
            pdf_path = os.path.join(output_dir, f"{resume_id}.pdf")
            pdf_success = builder.to_pdf(pdf_path)
            if pdf_success:
                results['output_files']['pdf'] = pdf_path

            results['resume_id'] = resume_id

            # Step 6: Track in database
            safe_print("\n[6/6] Tracking application in pgai...")

            job_title = job_analysis.get('job_title', 'Unknown Position')

            tracking_success = self.pgai.save_resume_application(
                company=company_name,
                job_title=job_title,
                job_posting=job_posting_text[:2000],  # Store excerpt
                resume_id=resume_id,
                keywords=job_analysis.get('raw_keywords', []),
                matched_skills=prioritized_skills,
                years_required=job_analysis.get('years_experience', 0)
            )

            results['tracking_record'] = tracking_success
            results['success'] = docx_success or pdf_success

            safe_print("\n" + "="*70)
            safe_print("RESUME GENERATION COMPLETE")
            safe_print("="*70)
            safe_print(f"Resume ID: {resume_id}")
            safe_print(f"Job: {job_title} at {company_name}")
            safe_print(f"Skills highlighted: {len(prioritized_skills)}")
            safe_print(f"Output files:")
            for fmt, path in results['output_files'].items():
                safe_print(f"  - {fmt.upper()}: {path}")

            return results

        except Exception as e:
            safe_print(f"ERROR: {e}")
            import traceback
            traceback.print_exc()
            return results

    def track_application(self, company: str, job_id: str, resume_id: str):
        """Track sent resume application"""
        safe_print(f"Tracking: {resume_id} -> {company}")

        try:
            cur = self.pgai.conn.cursor()

            cur.execute("""
                UPDATE resume_applications
                SET sent_date = NOW()
                WHERE resume_id = %s AND company_name = %s
            """, (resume_id, company))

            self.pgai.conn.commit()
            safe_print(f"  Tracked successfully")
            return True

        except Exception as e:
            safe_print(f"  Tracking error: {e}")
            return False

    def list_memories_for_query(self, query: str, limit: int = 10):
        """List memories matching a query (for debugging/testing)"""
        safe_print(f"\nMemories matching: '{query}'")
        safe_print("-" * 70)

        results = self.pgai.query_memories(query, limit=limit)

        for i, memory in enumerate(results, 1):
            safe_print(f"\n[{i}] Namespace: {memory.get('namespace', 'unknown')}")
            content = memory.get('content', '')[:150]
            safe_print(f"    Content: {content}...")

        return results

    def cleanup(self):
        """Close database connection"""
        if self.pgai:
            self.pgai.close()


# ============================================================================
# CLI INTERFACE
# ============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="ATS-Optimized Resume Generation Engine",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Generate resume from job posting text
  python resume_generator.py --mode generate \\
    --job-posting "Python Developer, 5+ years experience..." \\
    --company "TechCorp Inc" \\
    --output ./resumes

  # List memories for debugging
  python resume_generator.py --mode list-memories \\
    --query "Python experience"

  # Track sent application
  python resume_generator.py --mode track \\
    --company "TechCorp Inc" \\
    --resume-id resume_TechCorp_a1b2c3d4
        """
    )

    parser.add_argument(
        '--mode',
        choices=['generate', 'track', 'list-memories'],
        default='generate',
        help='Operation mode'
    )

    parser.add_argument(
        '--job-posting',
        type=str,
        help='Job posting text (can also read from file with @filename)'
    )

    parser.add_argument(
        '--company',
        type=str,
        default='Unknown Company',
        help='Company name'
    )

    parser.add_argument(
        '--output',
        type=str,
        default='./generated_resumes',
        help='Output directory for resume files'
    )

    parser.add_argument(
        '--resume-id',
        type=str,
        help='Resume ID for tracking'
    )

    parser.add_argument(
        '--query',
        type=str,
        help='Query for memory search'
    )

    args = parser.parse_args()

    # Initialize engine
    engine = ResumeGenerationEngine()

    try:
        if args.mode == 'generate':
            if not args.job_posting:
                parser.error("--job-posting required for generate mode")

            # Handle file input
            job_text = args.job_posting
            if job_text.startswith('@'):
                with open(job_text[1:], 'r') as f:
                    job_text = f.read()

            results = engine.generate(
                job_posting_text=job_text,
                output_dir=args.output,
                company_name=args.company
            )

            # Return exit code
            sys.exit(0 if results['success'] else 1)

        elif args.mode == 'track':
            if not args.resume_id:
                parser.error("--resume-id required for track mode")

            engine.track_application(args.company, '', args.resume_id)
            sys.exit(0)

        elif args.mode == 'list-memories':
            if not args.query:
                parser.error("--query required for list-memories mode")

            engine.list_memories_for_query(args.query, limit=10)
            sys.exit(0)

    finally:
        engine.cleanup()


if __name__ == '__main__':
    main()
